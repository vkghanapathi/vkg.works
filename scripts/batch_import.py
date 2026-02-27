"""
batch_import.py — One-time bulk import of VKG's composition archive into vkg-works.

Processes a ZIP or DOCX file:
  - Extracts DOCX + PDF files
  - Reads title from first heading/paragraph of each DOCX
  - Creates a .md sidecar with frontmatter (status: draft → appears in Queue)
  - Routes to the correct content/{section}/ directory
  - Preserves matching PDFs as companion download files

Usage:
    python scripts/batch_import.py --zip PATH --section articles [--category analytical]
    python scripts/batch_import.py --docx PATH --section articles [--category analytical]

All imported content gets status: draft so VKG can review in /queue/ before publishing.
"""
from __future__ import annotations

import io
import sys
# Force UTF-8 output on Windows (cp1252 can't print arrows / Sanskrit)
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
if hasattr(sys.stderr, 'reconfigure'):
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')

import argparse
import re
import sys
import tempfile
import zipfile
from pathlib import Path

ROOT = Path(__file__).parent.parent
CONTENT_DIR = ROOT / 'content'

AUTHOR = 'Dr. Vamshi Krishna Ghanapāṭhī'

# File extensions we can use as article content sources
CONTENT_EXTS = {'.docx'}
# File extensions we copy as companions (PDF download)
COMPANION_EXTS = {'.pdf'}
# File extensions we skip entirely
SKIP_EXTS = {'.xlsx', '.xls', '.xps', '.lp2', '.lp', '.mp3', '.m4a',
             '.doc', '.pptx', '.txt', '.wav', '.ogg'}


def _slugify(name: str) -> str:
    """Convert a filename stem to a URL-safe slug."""
    # Strip leading number prefix like "51 " or "P98 "
    name = re.sub(r'^[A-Za-z]?\d+\s+', '', name)
    name = name.lower()
    name = re.sub(r'[^\w\s-]', '', name)
    name = re.sub(r'[\s_]+', '-', name).strip('-')
    return name[:80]


def _extract_title(docx_path: Path, fallback: str) -> str:
    """Extract title from DOCX: first Heading, or first short paragraph."""
    try:
        import docx as _docx
        doc = _docx.Document(str(docx_path))
        # Prefer heading styles
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if 'heading' in para.style.name.lower() or para.style.name == 'Title':
                return text[:200]
        # Fall back to first non-empty paragraph if short enough to be a title
        for para in doc.paragraphs:
            text = para.text.strip()
            if text and len(text) <= 120:
                return text
        return fallback
    except Exception:
        return fallback


def _title_from_filename(stem: str) -> str:
    """Convert filename stem to a readable title (preserve original capitalisation)."""
    # Strip leading number prefix like "51 " or "R1" or "P98"
    title = re.sub(r'^[A-Za-z]?\d+\s*', '', stem).strip()
    return title if title else stem


def process_docx(docx_path: Path, section: str, category: str | None,
                 dest_dir: Path, pdf_companions: dict[str, Path]) -> bool:
    """
    Process a single DOCX file → write .md sidecar + copy DOCX + optional PDF.
    Returns True if file was processed, False if skipped.
    """
    if docx_path.stat().st_size == 0:
        print(f'  SKIP (empty): {docx_path.name}')
        return False

    stem = docx_path.stem
    slug = _slugify(stem)
    if not slug:
        slug = re.sub(r'[^\w-]', '-', stem.lower())[:60]

    title_fallback = _title_from_filename(stem)
    title = _extract_title(docx_path, title_fallback)

    date_str = '2026-02-27'
    dest_base = f'{date_str}-{slug}'
    dest_docx = dest_dir / f'{dest_base}.docx'

    if dest_docx.exists() and not force:
        # Still regenerate the .md sidecar in case it was malformed
        sidecar_only = dest_dir / f'{dest_base}.md'
        if not sidecar_only.exists():
            pass  # fall through to write sidecar
        else:
            print(f'  SKIP (exists): {dest_docx.name}')
            return False

    # Copy DOCX
    import shutil
    dest_dir.mkdir(parents=True, exist_ok=True)
    shutil.copy2(docx_path, dest_docx)

    # Copy PDF companion if available
    pdf_src = pdf_companions.get(stem)
    if pdf_src and pdf_src.exists():
        shutil.copy2(pdf_src, dest_dir / f'{dest_base}.pdf')

    # Write .md sidecar with frontmatter (use yaml.dump for safe quoting)
    import yaml as _yaml
    sidecar = dest_dir / f'{dest_base}.md'
    meta: dict = {
        'title': title,
        'date': date_str,
        'author': AUTHOR,
        'status': 'draft',
    }
    if category:
        meta['category'] = category
    yaml_block = _yaml.dump(meta, allow_unicode=True, default_flow_style=False,
                            sort_keys=False).rstrip()
    sidecar.write_text(f'---\n{yaml_block}\n---\n', encoding='utf-8')

    print(f'  OK: [{section}] {dest_base}  — "{title}"')
    return True


def ingest_zip(zip_path: Path, section: str, category: str | None) -> tuple[int, int]:
    """Extract ZIP, process DOCX files into content/{section}/."""
    dest_dir = CONTENT_DIR / section
    ingested = skipped = 0

    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        with zipfile.ZipFile(zip_path, 'r') as zf:
            # Extract all members (ignore encoding errors in filenames)
            for member in zf.infolist():
                try:
                    zf.extract(member, tmp_path)
                except Exception as e:
                    print(f'  WARN: could not extract {member.filename}: {e}')

        # Build index of PDF companions keyed by stem
        pdf_companions: dict[str, Path] = {}
        for pdf in tmp_path.rglob('*.pdf'):
            pdf_companions[pdf.stem] = pdf

        # Process DOCX files
        for docx in sorted(tmp_path.rglob('*.docx')):
            ext = docx.suffix.lower()
            if ext not in CONTENT_EXTS:
                continue
            result = process_docx(docx, section, category, dest_dir, pdf_companions)
            if result:
                ingested += 1
            else:
                skipped += 1

    return ingested, skipped


def ingest_docx(docx_path: Path, section: str, category: str | None) -> bool:
    """Process a standalone DOCX file."""
    dest_dir = CONTENT_DIR / section
    # Check for sibling PDF
    pdf_path = docx_path.with_suffix('.pdf')
    companions = {docx_path.stem: pdf_path} if pdf_path.exists() else {}
    return process_docx(docx_path, section, category, dest_dir, companions)


def main() -> None:
    parser = argparse.ArgumentParser(description='Bulk import DOCX/ZIP into vkg-works content/')
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument('--zip', help='Path to ZIP archive')
    group.add_argument('--docx', help='Path to standalone DOCX file')
    parser.add_argument('--section', required=True,
                        choices=['articles', 'poems', 'songs', 'books', 'audio',
                                 'video', 'projects', 'coverage'],
                        help='Target content section')
    parser.add_argument('--category', default=None,
                        help='Category tag (e.g. analytical, administrative)')
    args = parser.parse_args()

    if args.zip:
        zip_path = Path(args.zip)
        if not zip_path.exists():
            print(f'ERROR: ZIP not found: {zip_path}', file=sys.stderr)
            sys.exit(1)
        print(f'Processing ZIP: {zip_path.name} -> content/{args.section}/')
        ingested, skipped = ingest_zip(zip_path, args.section, args.category)
        print(f'\nDone: {ingested} imported, {skipped} skipped.')
    else:
        docx_path = Path(args.docx)
        if not docx_path.exists():
            print(f'ERROR: DOCX not found: {docx_path}', file=sys.stderr)
            sys.exit(1)
        print(f'Processing DOCX: {docx_path.name} -> content/{args.section}/')
        result = ingest_docx(docx_path, args.section, args.category)
        print('Done.' if result else 'Skipped.')


if __name__ == '__main__':
    main()
