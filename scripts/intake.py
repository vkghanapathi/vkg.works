"""
intake.py — Unified batch intake pipeline for vkg.works.

Chains in one command:
  1. Extract DOCX/PDF files from a ZIP (or process a single DOCX/PDF)
  2. Read full document text
  3. Call Gemini API (via gcloud ADC) → infer title, language, subject, topic,
     abstract, preamble, keywords — all in a SINGLE API call per document
  4. Write complete .md sidecar with ALL metadata fields populated
  5. Assign VKG-X-### UID from state/registry.json
  6. status: draft  →  item appears in /queue/ for VKG review before publishing

Usage:
    python scripts/intake.py --zip inbox/batch.zip --section articles
    python scripts/intake.py --docx path/to/essay.docx --section poems
    python scripts/intake.py --pdf  path/to/book.pdf  --section books
    python scripts/intake.py --zip batch.zip --section songs --date 2026-03-01
    python scripts/intake.py --zip batch.zip --section articles --dry-run

Requires:
    gcloud CLI authenticated (gcloud auth login / application-default login)
    pip install python-frontmatter python-docx pyyaml pdfplumber
"""
from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
import tempfile
import time
import zipfile
from datetime import date
from pathlib import Path

import frontmatter
import yaml

# Gemini client (uses gcloud ADC — no API key file needed)
sys.path.insert(0, str(Path(__file__).parent))
from _gemini_rest import GeminiClient

# ── Force UTF-8 output on Windows ────────────────────────────────────────────
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
if hasattr(sys.stderr, 'reconfigure'):
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')

ROOT         = Path(__file__).parent.parent
CONTENT_DIR  = ROOT / 'content'
REGISTRY_FILE = ROOT / 'state' / 'registry.json'

AUTHOR       = 'Dr. Vamshi Krishna Ghanapāṭhī'
AUTHOR_ORCID = '0009-0007-3852-0158'
API_DELAY    = 1.0   # seconds between calls

SECTION_PREFIX: dict[str, str] = {
    'articles': 'A', 'poems': 'P', 'songs': 'S', 'books': 'B',
    'audio': 'AU', 'video': 'V', 'projects': 'PR', 'coverage': 'C',
}

KNOWN_SECTIONS = set(SECTION_PREFIX.keys())

SUBJECT_CHOICES = [
    'Vedic Ritual',
    'Devotional Music',
    'Philosophy & Vedānta',
    'Sacred Poetry',
    'Dharmaśāstra',
    'Jyotiṣa',
    'Vedic Linguistics',
    'Contemporary Commentary',
]

BODY_LIMIT = 4000   # characters sent to Gemini

# ── Prompts ───────────────────────────────────────────────────────────────────

SYSTEM_PROMPT = """\
You are a scholarly metadata assistant specialising in Vedic literature,
Sanskrit, Telugu, and Kannada compositions by Dr. Vamshi Krishna Ghanapāṭhī.
You infer accurate metadata from document text without translating or
paraphrasing the original. All descriptive output is in English.
"""

PARSE_PROMPT = """\
Document text (first {char_limit} characters):
---
{body}
---
Filename hint : {filename}
Section hint  : {section}

Infer the following metadata for this work.
Respond with ONLY a valid JSON object — no markdown, no explanation:

{{
  "title": "<Full title of the work — infer from heading or first line. Clean, readable.>",
  "language": "<Primary script/language — ISO 639-1: sa (Sanskrit), te (Telugu), kn (Kannada), en (English). Use semicolons for mixed: sa;te>",
  "subject": "<One subject from this list only: {subjects}>",
  "topic": ["<specific term>", "<specific term>", "<specific term>"],
  "abstract": "<2–3 sentences: what this work is and its scholarly significance.>",
  "preamble": "<3–5 sentences: scriptural/ritual/devotional context for a reader new to the tradition.>",
  "keywords": ["<term1>", "<term2>", "<term3>", "<term4>", "<term5>"]
}}

Rules:
- title: extract from first heading or short first line; do NOT invent
- language: reflects the composition language (not metadata). Devanagari → sa; Telugu script → te
- subject: pick EXACTLY ONE from the list
- topic: 2–4 specific names (deity, scripture, ritual, concept)
- abstract: scholarly, what the text IS — do not just restate the title
- preamble: tradition, scripture, deity, or practice named — for a general reader
- keywords: 5–8 specific subject terms (deity names, scripture names, ritual names, concepts)
"""


# ── Utilities ─────────────────────────────────────────────────────────────────

def _slugify(name: str) -> str:
    name = re.sub(r'^[A-Za-z]?\d+\s*', '', name)   # strip leading number prefix
    name = name.lower()
    name = re.sub(r'[^\w\s-]', '', name)
    name = re.sub(r'[\s_]+', '-', name).strip('-')
    return name[:80] or 'untitled'


def _extract_docx_text(docx_path: Path, limit: int = BODY_LIMIT) -> str:
    try:
        from docx import Document
        doc = Document(str(docx_path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        return '\n'.join(paragraphs)[:limit]
    except Exception as e:
        print(f'    WARN: could not read DOCX text ({e})', file=sys.stderr)
        return ''


def _extract_pdf_text(pdf_path: Path, limit: int = BODY_LIMIT) -> str:
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(str(pdf_path)) as pdf:
            for page in pdf.pages[:12]:   # first 12 pages is enough for context
                t = page.extract_text()
                if t:
                    text_parts.append(t.strip())
                if sum(len(p) for p in text_parts) >= limit:
                    break
        return '\n'.join(text_parts)[:limit]
    except Exception as e:
        print(f'    WARN: could not read PDF text ({e})', file=sys.stderr)
        return ''


def _extract_text(path: Path, limit: int = BODY_LIMIT) -> str:
    """Extract text from DOCX or PDF."""
    if path.suffix.lower() == '.pdf':
        return _extract_pdf_text(path, limit)
    return _extract_docx_text(path, limit)


def _extract_title_from_docx(docx_path: Path) -> str:
    """Quick title extraction — fallback before Claude runs."""
    try:
        from docx import Document
        doc = Document(str(docx_path))
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if 'heading' in para.style.name.lower() or para.style.name == 'Title':
                return text[:200]
        for para in doc.paragraphs:
            text = para.text.strip()
            if text and len(text) <= 120:
                return text
    except Exception:
        pass
    return ''


def _next_uid(section: str, registry: dict) -> tuple[str, int]:
    """Allocate next UID for section; returns (uid_string, int_number)."""
    prefix = SECTION_PREFIX.get(section, section.upper()[:2])
    items: dict = registry.setdefault(section, {})
    next_num = (max(items.values(), default=0) + 1)
    return f'VKG-{prefix}-{next_num:03d}', next_num


def _save_registry(registry: dict) -> None:
    REGISTRY_FILE.write_text(
        json.dumps(registry, indent=2, ensure_ascii=False),
        encoding='utf-8'
    )


# ── Gemini API call ───────────────────────────────────────────────────────────

def _call_gemini(body: str, filename: str, section: str,
                 client: GeminiClient) -> dict | None:
    prompt = PARSE_PROMPT.format(
        char_limit=BODY_LIMIT,
        body=body or '(text could not be extracted)',
        filename=filename,
        section=section,
        subjects=', '.join(SUBJECT_CHOICES),
    )
    return client.generate_json(SYSTEM_PROMPT, prompt, max_tokens=1024)


# ── Sidecar writer ────────────────────────────────────────────────────────────

def _write_sidecar(
    dest_dir: Path,
    dest_base: str,
    meta: dict,
    dry_run: bool,
) -> Path:
    sidecar = dest_dir / f'{dest_base}.md'
    yaml_block = yaml.dump(
        meta, allow_unicode=True, default_flow_style=False, sort_keys=False
    ).rstrip()
    content = f'---\n{yaml_block}\n---\n'
    if not dry_run:
        dest_dir.mkdir(parents=True, exist_ok=True)
        sidecar.write_text(content, encoding='utf-8')
    return sidecar


# ── Process one DOCX ─────────────────────────────────────────────────────────

def process_one(
    docx_path: Path,
    section: str,
    date_str: str,
    client: GeminiClient,
    registry: dict,
    pdf_companions: dict[str, Path],
    dry_run: bool,
    verbose: bool = True,
) -> bool:
    """
    Full intake pipeline for a single DOCX/PDF file.
    Returns True on success.
    """
    stem = docx_path.stem

    # Quick fallback title before AI call
    quick_title = _extract_title_from_docx(docx_path) or stem
    slug = _slugify(stem)
    dest_base = f'{date_str}-{slug}'
    dest_dir  = CONTENT_DIR / section

    # Check for duplicate
    if (dest_dir / f'{dest_base}.md').exists() and not dry_run:
        print(f'  SKIP (exists): {dest_base}.md')
        return False

    # Extract text and call Gemini
    body = _extract_text(docx_path)
    print(f'  → Calling Gemini for: {stem[:60]}')
    parsed = _call_gemini(body, stem, section, client)
    time.sleep(API_DELAY)

    if parsed is None:
        print(f'  WARN: Gemini returned no data — using filename as title')
        parsed = {}

    # Build metadata — Claude fills most fields, we fill the rest
    title     = str(parsed.get('title', '') or quick_title).strip()
    language  = str(parsed.get('language', '')).strip()
    subject   = str(parsed.get('subject', '')).strip()
    topic     = parsed.get('topic', [])
    abstract  = str(parsed.get('abstract', '')).strip()
    preamble  = str(parsed.get('preamble', '')).strip()
    keywords  = parsed.get('keywords', [])

    if isinstance(topic,    str): topic    = [topic]
    if isinstance(keywords, str): keywords = [keywords]
    topic    = [str(t).strip() for t in topic    if t]
    keywords = [str(k).strip() for k in keywords if k]

    # Allocate UID
    uid, uid_num = _next_uid(section, registry)

    meta: dict = {
        'title':    title,
        'date':     date_str,
        'author':   AUTHOR,
        'uid':      uid,
        'orcid':    AUTHOR_ORCID,
        'doi':      '',
        'status':   'draft',
        'section':  section,
    }
    if language: meta['language'] = language
    if subject:  meta['subject']  = subject
    if topic:    meta['topic']    = topic
    if abstract: meta['abstract'] = abstract
    if preamble: meta['preamble'] = preamble
    if keywords: meta['keywords'] = keywords

    # Register slug → uid_num
    slug_key = re.sub(r'^\d{4}-\d{2}-\d{2}-', '', dest_base)  # bare slug
    if not dry_run:
        registry[section][slug_key] = uid_num

    # Write sidecar
    _write_sidecar(dest_dir, dest_base, meta, dry_run)

    # Copy source file and PDF companion
    if not dry_run:
        dest_dir.mkdir(parents=True, exist_ok=True)
        ext = docx_path.suffix.lower()
        shutil.copy2(docx_path, dest_dir / f'{dest_base}{ext}')
        if ext != '.pdf':
            pdf_src = pdf_companions.get(stem)
            if pdf_src and pdf_src.exists():
                shutil.copy2(pdf_src, dest_dir / f'{dest_base}.pdf')

    tag = 'DRY' if dry_run else 'OK'
    print(f'  {tag}: [{section}] {uid}  {dest_base}')
    print(f'       Title    : {title}')
    print(f'       Language : {language or "—"}  |  Subject: {subject or "—"}')
    print(f'       Topics   : {", ".join(topic) if topic else "—"}')
    return True


# ── ZIP handler ───────────────────────────────────────────────────────────────

def ingest_zip(
    zip_path: Path,
    section: str,
    date_str: str,
    client: GeminiClient,
    registry: dict,
    dry_run: bool,
) -> tuple[int, int]:
    ingested = skipped = 0
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        with zipfile.ZipFile(zip_path, 'r') as zf:
            for member in zf.infolist():
                try:
                    zf.extract(member, tmp_path)
                except Exception as e:
                    print(f'  WARN: could not extract {member.filename}: {e}')

        # Index PDF companions by stem
        pdf_companions: dict[str, Path] = {
            p.stem: p for p in tmp_path.rglob('*.pdf')
        }

        source_files = sorted(
            [f for f in tmp_path.rglob('*')
             if f.suffix.lower() in ('.docx', '.pdf')
             and not f.name.startswith(('.', '~'))]
        )
        # Don't double-process PDFs that are companions to a DOCX
        docx_stems = {f.stem for f in source_files if f.suffix.lower() == '.docx'}
        source_files = [
            f for f in source_files
            if not (f.suffix.lower() == '.pdf' and f.stem in docx_stems)
        ]
        if not source_files:
            print('  WARN: no DOCX or PDF files found in ZIP')
            return 0, 0

        for docx in source_files:
            ok = process_one(docx, section, date_str, client,
                             registry, pdf_companions, dry_run)
            if ok:
                ingested += 1
            else:
                skipped += 1

    return ingested, skipped


# ── Main ──────────────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(
        description='Intake DOCX/ZIP → AI metadata → .md sidecar → /queue/'
    )
    src = parser.add_mutually_exclusive_group(required=True)
    src.add_argument('--zip',  type=Path, metavar='PATH', help='ZIP archive of DOCX/PDF files')
    src.add_argument('--docx', type=Path, metavar='PATH', help='Single DOCX file')
    src.add_argument('--pdf',  type=Path, metavar='PATH', help='Single PDF file')

    parser.add_argument('--section', required=True, choices=sorted(KNOWN_SECTIONS),
                        help='Target section (articles, poems, songs…)')
    parser.add_argument('--date', default=str(date.today()),
                        metavar='YYYY-MM-DD', help='Publication date (default: today)')
    parser.add_argument('--dry-run', action='store_true',
                        help='Preview without writing files or updating registry')
    args = parser.parse_args()

    # Create Gemini client (uses gcloud ADC — no API key needed)
    client = GeminiClient()

    # Load registry
    if REGISTRY_FILE.exists():
        with open(REGISTRY_FILE, encoding='utf-8') as f:
            registry = json.load(f)
    else:
        registry = {}

    print(f'\n── VKG Intake Pipeline ──────────────────────────────────────')
    print(f'  Section  : {args.section}')
    print(f'  Date     : {args.date}')
    print(f'  Engine   : Gemini 2.0 Flash (gcloud ADC)')
    print(f'  Dry run  : {args.dry_run}')
    print()

    ingested = skipped = 0

    if args.zip:
        if not args.zip.exists():
            print(f'ERROR: ZIP not found: {args.zip}', file=sys.stderr)
            sys.exit(1)
        ingested, skipped = ingest_zip(
            args.zip, args.section, args.date,
            client, registry, args.dry_run,
        )
    else:
        single = args.docx or args.pdf
        if not single.exists():
            print(f'ERROR: File not found: {single}', file=sys.stderr)
            sys.exit(1)
        ok = process_one(
            single, args.section, args.date,
            client, registry,
            pdf_companions={}, dry_run=args.dry_run,
        )
        ingested = 1 if ok else 0
        skipped  = 0 if ok else 1

    # Save updated registry
    if not args.dry_run and ingested > 0:
        _save_registry(registry)
        print(f'\n  Registry updated → state/registry.json')

    print(f'\n── Done ─────────────────────────────────────────────────────')
    print(f'  Ingested : {ingested}')
    print(f'  Skipped  : {skipped}')
    if ingested > 0 and not args.dry_run:
        print(f'\n  Items are now in /queue/ (status: draft).')
        print(f'  Review at vkg.works/queue/ → change status to published → push.')


if __name__ == '__main__':
    main()
